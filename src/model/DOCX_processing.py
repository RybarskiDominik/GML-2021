from PySide6.QtWidgets import QApplication, QFileDialog
from docx.shared import Inches
from docx import Document
from model.python_docx_replace import docx_replace
import sys, os

def inwestigation_docx(docx_path, output_path, data):
    doc = Document(docx_path)
    docx_replace(doc, data)
    doc.save(output_path)

def inwestigation_docx_OLD(docx_path, output_path, data):
    
    doc = Document(docx_path)
    xml_content = doc._element
    namespaces = doc.element.nsmap
    
    def replace_text_in_runs_simple(paragraph, data):
        for run in paragraph.runs:
            for key, value in data.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    def replace_text_in_runs(paragraph, data):
        full_text = "".join(run.text for run in paragraph.runs)  # Merging divided tags
        """ 
        full_text = ""
        for run in paragraph.runs:
            full_text += run.text
        """
        modified = False  # Flag pointing whether a swap was made

        for key, value in data.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
                modified = True

        if modified:  # If a replacement is made, we overwrite the text in the paragraph
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(full_text)

    # Replacement in regular paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_runs_simple(paragraph, data)

    # Replacement in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_runs(paragraph, data)
    
    # Replacement in header and footer
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs: #header = section.header, footer = section.footer
            for run in paragraph.runs:
                for key, value in data.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)


    # Replacement in textboxes
    for shape in xml_content.xpath("//w:txbxContent"):  # Replace text in shapes (textboxes) using XML
        for node in shape.xpath(".//w:t", namespaces=namespaces):

            if node.text:
                for key, value in data.items():
                    if key in node.text:
                        node.text = node.text.replace(key, value)

    """
    # Replacement in header and footer
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs: #header = section.header, footer = section.footer
            replace_text_in_runs(paragraph, data)
    
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for paragraph in header_footer.paragraphs:
                if not any(run.element.tag == 'w:drawing' for run in paragraph.runs): 
                    replace_text_in_runs(paragraph, data)

    for shape in doc.inline_shapes:
        if shape._element.xpath(".//w:t"):  # Sprawdzenie, czy kształt zawiera tekst
            for node in shape._element.xpath(".//w:t"):
                for key, value in data.items():
                    if node.text and key in node.text:
                        node.text = node.text.replace(key, value)
    """

    doc.save(output_path)

def select_folder_and_process(path_to_docx=None, output_folder=None, data=None):

    if data is None:
        data = {}

    if not path_to_docx:
        path_to_docx = QFileDialog.getExistingDirectory(None, "Select Folder with DOCX Files")  # Select input folder
    if not path_to_docx:
        return 0

    output_folder = QFileDialog.getExistingDirectory(None, "Select Output Folder", os.path.expanduser("~/Desktop"))  # Select output folder
    if not output_folder:
        return 0
    
    if os.path.abspath(path_to_docx) == os.path.abspath(output_folder):
        return 0

    for filename in os.listdir(path_to_docx):  # Process all files in the selected folder
        if filename.endswith('.docx'):
            docx_path = os.path.join(path_to_docx, filename)
            output_path = os.path.join(output_folder, filename)  # f"modified_{filename}"
            inwestigation_docx(docx_path, output_path, data)
    return 1


if __name__ == "__main__":
    app = QApplication(sys.argv)

    path = os.path.join(sys.path[0], 'test.docx')
    path_out = os.path.join(sys.path[0], 'Out.docx')

    data = {'[NR_DZIALKI]': 'A',
            '[IDENTYFIKATOR]': 'B',
            '[WOJ]': 'C',
            '[POW]': 'D',
            '[JEWID]': 'E',
            '[JEWID_ID]': 'F',
            '[OBR]': 'G',
            '[OBR_ID]': 'H',
            '[ARKUSZ]': 'I',
            '[DATA_OKLADKA]': 'J',
            '[DATA_SPIS]': 'K',
            '[DATA_SPR]': 'L',
            '[DATA]': 'M',
            '[CEL]': 'N',
            '[WYKONAWCA]': 'O',
            '[KIEROWNIK]': 'P',
            '[KIEROWNIK_UPR]': 'R',
            '[UPRAC]': 'S',
            '[TERMIN_R]': 'T',
            '[TERMIN_Z]': 'U',
            }
    
    inwestigation_docx(path, path_out, data)

    #select_folder_and_process(data=data)

